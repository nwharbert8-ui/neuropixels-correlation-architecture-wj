# -*- coding: utf-8 -*-
"""Assemble the Journal of Neurophysiology submission folder for the Neuropixels paper.
Author: Drake H. Harbert (D.H.H.) | ORCID 0009-0007-7740-3616 | 2026-06-30"""
import os, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
OUT = r"G:/My Drive/inner_architecture_research/neuropixels_wj/JNeurophysiol_Submission"
TITLE = ("Population correlation architecture in mouse visual cortex reorganizes by "
         "magnitude, not by sign, across visual stimulus conditions")
C = WD_ALIGN_PARAGRAPH.CENTER
def newdoc():
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Inches(1.0); s.left_margin = s.right_margin = Inches(1.0)
    d.styles["Normal"].font.name = "Times New Roman"; d.styles["Normal"].font.size = Pt(12)
    return d
def P(d, t, bold=False, italic=False, align=None, after=8, sp=1.15, size=12):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = sp
    if align: p.alignment = align
    r = p.add_run(t); r.bold = bold; r.italic = italic; r.font.size = Pt(size); return p

# ---- Title page ----
d = newdoc()
P(d, TITLE, bold=True, align=C, after=14, sp=2.0, size=14)
P(d, "Drake H. Harbert", align=C, after=2)
P(d, "Inner Architecture LLC, Canton, OH 44720, United States", align=C, italic=True, size=11, after=2)
P(d, "ORCID: 0009-0007-7740-3616", align=C, size=11, after=2)
P(d, "Correspondence: Drake H. Harbert, Drake@InnerArchitectureLLC.com", align=C, size=11, after=14)
P(d, "Running head: Magnitude-driven reorganization in visual cortex", after=2)
P(d, "Keywords: population coding; spike-count correlations; Neuropixels; visual cortex; "
      "weighted Jaccard; correlation architecture; sign preservation", after=2)
P(d, "Conflicts of interest: The author is the founder of Inner Architecture LLC and declares "
      "no other competing interests.", after=2)
P(d, "Funding: This research received no external funding.", after=2)
d.save(os.path.join(OUT, "Title_Page.docx"))

# ---- Cover letter ----
d = newdoc()
P(d, "Drake H. Harbert", after=0); P(d, "Inner Architecture LLC, Canton, OH, United States", size=10, after=0)
P(d, "Drake@InnerArchitectureLLC.com  |  ORCID 0009-0007-7740-3616", size=10, after=12)
P(d, "June 30, 2026", after=0)
P(d, "The Editors, Journal of Neurophysiology (American Physiological Society)", after=12)
P(d, "Re: Submission of an original research article", bold=True, after=10)
P(d, "Dear Editors,")
for t in [
 "Please consider the enclosed manuscript, \"" + TITLE + ",\" as an original research article for the Journal of Neurophysiology.",
 "Pairwise spike-count correlations among neurons are central to population coding, yet they are usually summarized by a scalar or compared by unsigned similarity, which treats a correlation sign reversal as no change. Using 30 sessions of the Allen Brain Observatory Visual Coding Neuropixels dataset (2,052 condition comparisons, ~55,000 units), we measured sign reversal directly and stratified by magnitude. Correlation architecture reorganizes substantially between visual conditions, but at every meaningful correlation magnitude the change is in strength, not sign: sign is preserved in more than 98% of pairs, and reorganization is a graded population property only weakly related to spike waveform. A scalar summary changes roughly 60-fold less than the architecture over the same comparisons, and each neuron expresses a small, structured set of coupling regimes.",
 "We believe the work fits the journal's scope in systems and population neurophysiology, and that the direct, magnitude-stratified treatment of correlation sign is a methodological point of general use. Inference treats sessions, not non-independent pairs, as the replication unit where appropriate, and the limitations (correlational design, somatic integration of dendritic inputs, the relational nature of the coupling-regime measure) are stated explicitly.",
 "We intend to publish under the standard subscription option at no charge. The manuscript is original, is not under consideration elsewhere, and the sole author has approved this submission. All data are public (DANDI 000021) and all analysis code is openly available.",
 "Thank you for your consideration.",
 "Respectfully,", "Drake H. Harbert",
]: P(d, t)
d.save(os.path.join(OUT, "Cover_Letter.docx"))

# ---- Figure captions ----
d = newdoc()
P(d, "Figure Legends", bold=True, after=10, size=13)
caps = [
 ("Figure 1.", "Direct, magnitude-stratified sign-flip rate across 2,052 condition comparisons "
  "(30 sessions). The fraction of neuron pairs reversing correlation sign between conditions is "
  "shown as a function of the magnitude stratum min(|r_A|, |r_B|). The rate is below the 50% chance "
  "line at every stratum and falls to ~1% at |r| >= 0.2 (every comparison below chance, binomial "
  "p ~ 0). Reorganization is in correlation magnitude, not sign."),
 ("Figure 2.", "Per-neuron row reorganization versus sign coherence (hexbin density, ~55,000 "
  "neurons). Neurons that restructure their coupling more strongly preserve sign less (Spearman "
  "rho = -0.52, same sign in all 30 of 30 sessions). The relationship is continuous; the split-"
  "learner neurons are the tail of this trade-off, not a discrete class."),
 ("Figure 3.", "What a scalar summary discards. Distributions, across 2,052 condition comparisons, "
  "of the change in the field-standard scalar (mean absolute correlation; median 0.009) and of the "
  "architecture reorganization (1 - unsigned weighted Jaccard; median 0.58). The scalar compresses "
  "the reorganization roughly 60-fold."),
 ("Figure 4.", "Per-neuron effective coupling regimes (participation ratio of a neuron's coupling "
  "profile across conditions). Observed values (median ~4) are far below a condition-shuffled null "
  "(median ~11); 100% of neurons fall below the null median, indicating a small, structured coupling "
  "repertoire reused across visual conditions."),
]
for h, t in caps:
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.5
    r = p.add_run(h + " "); r.bold = True; r.font.size = Pt(12)
    r2 = p.add_run(t); r2.font.size = Pt(12)
d.save(os.path.join(OUT, "Figure_Captions.docx"))

# ---- reference verification ----
ref = """# Reference verification, Neuropixels manuscript (Journal of Neurophysiology)
# Generated 2026-06-30. All DOIs verified by independent web lookup this session.

| # | Reference | DOI | Status |
|---|-----------|-----|--------|
| 1 | Cohen MR, Kohn A. Measuring and interpreting neuronal correlations. Nat Neurosci 14:811-819, 2011 | 10.1038/nn.2842 | VERIFIED (nature.com; PMID 21709677) |
| 2 | Harbert DH. Sigma-1 and sigma-2 receptors exhibit divergent genome-wide co-expression architectures... Front Pharmacol 17:1830847, 2026 | 10.3389/fphar.2026.1830847 | VERIFIED (author's own published paper) |
| 3 | Niell CM, Stryker MP. Highly selective receptive fields in mouse visual cortex. J Neurosci 28:7520-7536, 2008 | 10.1523/JNEUROSCI.0623-08.2008 | VERIFIED (jneurosci.org; PMID 18650330) |
| 4 | Siegle JH, Jia X, Durand S, et al. Survey of spiking in the mouse visual system reveals functional hierarchy. Nature 592:86-92, 2021 | 10.1038/s41586-020-03171-x | VERIFIED (nature.com; the canonical Allen Visual Coding Neuropixels dataset paper) |

No fabricated or unresolvable references. New manuscript, new venue: no prior-version reference diff applies.
"""
open(os.path.join(OUT, "_reference_verification.md"), "w", encoding="utf-8").write(ref)

# ---- manifest ----
man = """# Journal of Neurophysiology submission package - Neuropixels (2026-06-30)
Venue: Journal of Neurophysiology (APS). FREE to publish via Subscribe-to-Open (no APC).

Upload:
1. Manuscript.docx  (title, New & Noteworthy, abstract, intro, methods, results, discussion, declarations, references)
2. Title_Page.docx
3. Cover_Letter.docx
4. Figure_Captions.docx
5. Main_Figures/Fig1-4 (PDF for submission; PNG provided)
6. Reference check: _reference_verification.md (all 4 DOIs verified)

Data: Allen Visual Coding Neuropixels (DANDI 000021). Code: github.com/nwharbert8-ui/neuropixels-correlation-architecture-wj (pushed).

Confirm on the APS portal at submission: exact abstract word limit, New & Noteworthy limit (~75 words),
reference style, and select the standard (no-fee, Subscribe-to-Open) option.
"""
open(os.path.join(OUT, "MANIFEST.md"), "w", encoding="utf-8").write(man)
print("Package assembled in", OUT)
print(sorted(os.listdir(OUT)))

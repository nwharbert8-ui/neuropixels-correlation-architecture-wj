# -*- coding: utf-8 -*-
"""
Build the brand-new Neuropixels manuscript for Journal of Neurophysiology (APS).
Author: Drake H. Harbert (D.H.H.) | ORCID 0009-0007-7740-3616 | 2026-06-30
Reports the corrected, tested results only. Conservative language. Single substrate.
"""
import os, sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
OUT = r"G:/My Drive/inner_architecture_research/neuropixels_wj/JNeurophysiol_Submission"
os.makedirs(OUT, exist_ok=True)
TITLE = ("Correlation-architecture reorganization in mouse visual cortex is predominantly "
         "in magnitude, with a small but genuine component of sign reversal above the "
         "sampling floor")

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1.0); s.left_margin = s.right_margin = Inches(1.0)
doc.styles["Normal"].font.name = "Times New Roman"; doc.styles["Normal"].font.size = Pt(12)
def H(t, size=13):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(size); return p
def P(t, after=8, sp=2.0, align=None, italic=False, bold=False, size=12):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = sp
    if align: p.alignment = align
    r = p.add_run(t); r.italic = italic; r.bold = bold; r.font.size = Pt(size); return p
C = WD_ALIGN_PARAGRAPH.CENTER

P(TITLE, bold=True, align=C, size=14, after=10)
P("Drake H. Harbert", align=C, after=2, sp=1.0)
P("Inner Architecture LLC, Canton, OH, United States", align=C, italic=True, size=11, after=2, sp=1.0)
P("ORCID: 0009-0007-7740-3616  |  Drake@InnerArchitectureLLC.com", align=C, size=10, after=12, sp=1.0)

H("New & Noteworthy")
P("Population analyses usually summarize pairwise correlations with their absolute "
  "magnitude, treating a sign reversal as no change. Applying a sign-resolved, "
  "threshold-free decomposition to 30 Neuropixels recordings and testing the sign-flip "
  "rate against a within-condition (sampling) null rather than a chance baseline, we show "
  "that reorganization between visual conditions is predominantly in magnitude, yet sign "
  "reversal, though rare (about 1% of strongly correlated pairs), significantly exceeds "
  "the sampling floor at every correlation magnitude and is thus a small but genuine "
  "component of the reorganization. Reorganization is a graded population property, not "
  "the property of a distinct cell class.", after=10)

H("Abstract")
P("Neural population activity is commonly characterized through the pairwise correlation "
  "structure among simultaneously recorded neurons, yet that structure is usually reduced "
  "to a single scalar or compared using unsigned similarity, which treats a correlation "
  "sign reversal as no change. We analyzed spike-sorted single-unit activity from 30 "
  "sessions of the Allen Brain Observatory Visual Coding Neuropixels dataset (1,305 to "
  "2,395 units per session, 6 to 8 visual areas), computing the full unit-by-unit "
  "Spearman correlation matrix within each visual stimulus condition and comparing matrices "
  "between conditions across 2,052 condition pairs. Reorganization between conditions was "
  "large by unsigned weighted Jaccard, and it was predominantly in magnitude: the direct, "
  "magnitude-stratified rate at which individual pairs reversed correlation sign fell to "
  "approximately 1% among pairs with correlation magnitude at least 0.2. Tested against a "
  "within-condition split-half (sampling) null rather than a chance baseline, however, the "
  "observed sign-flip rate significantly exceeded the sampling floor at every stratum (all "
  "30 sessions, Wilcoxon p < 0.001; at |r| >= 0.2, 1.2% observed vs 0.01% expected from "
  "sampling), so sign reversal is rare but genuine and its excess over the floor grows with "
  "correlation magnitude. A per-unit decomposition showed that most "
  "neurons restructured their coupling in strength while preserving sign (coherent-"
  "magnitude), a minority of about 4% split their partners by sign, and essentially none "
  "were unstructured. Reorganization traded off continuously against sign preservation "
  "(rho = -0.52, 30 of 30 sessions) and was only weakly related to extracellular waveform "
  "shape, indicating a graded population property rather than a cell-class signature. A "
  "field-standard scalar (mean correlation) changed roughly 60-fold less than the "
  "architecture between the same conditions. Each neuron expressed a small, structured set "
  "of coupling regimes across conditions (median ~4 of a possible ~12; below a condition-"
  "shuffled null in 100% of neurons). These results characterize how visual population "
  "correlation architecture is reorganized in a way single-metric analyses do not resolve.", after=10)
P("Keywords: population coding; spike-count correlations; Neuropixels; visual cortex; "
  "weighted Jaccard; correlation architecture; sign reversal", italic=True, size=11, after=10)

doc.add_page_break()
H("INTRODUCTION")
for t in [
 "The joint activity of a neural population is not captured by single-neuron responses "
 "alone; the pairwise correlations among neurons carry information about shared input, "
 "connectivity, and state. A large literature characterizes these spike-count (noise) "
 "correlations and how they depend on stimulus, attention, and arousal (Cohen and Kohn "
 "2011). Most analyses, however, reduce the correlation structure either to a scalar "
 "summary (a mean correlation) or to an unsigned comparison of correlation matrices, which "
 "treats a reversal from positive to negative coupling, the largest possible change in a "
 "relationship, as no change at all.",
 "Here we treat the correlation structure as an architecture and compare it between "
 "conditions with a sign-resolved, threshold-free decomposition. The fundamental unit is "
 "the individual spike-sorted neuron; the architecture is the full set of pairwise Spearman "
 "correlations among the simultaneously recorded population within a visual stimulus "
 "condition; and reorganization is the change in that architecture between conditions. We "
 "ask three questions the standard summaries cannot answer: how much of the reorganization "
 "is a change in correlation magnitude versus a reversal of correlation sign; whether any "
 "reorganization is concentrated in a distinguishable subset of neurons; and what a "
 "single-scalar summary discards relative to the full architecture.",
 "Using 30 sessions of the Allen Brain Observatory Visual Coding Neuropixels dataset "
 "(Siegle et al. 2021), we find that correlation architecture reorganizes substantially "
 "between visual conditions but that the reorganization is, at every meaningful correlation "
 "magnitude, a change in strength rather than in sign. Sign is preserved; magnitude is "
 "rearranged. The effect is graded across the population, only weakly related to spike "
 "waveform, and invisible to a scalar summary.",
]: P(t)

doc.add_page_break()
H("MATERIALS AND METHODS")
H("Dataset and fundamental unit", 12)
P("We used 30 sessions from the Allen Brain Observatory Visual Coding Neuropixels dataset "
  "(Siegle et al. 2021; DANDI 000021). The fundamental unit of analysis was the individual "
  "spike-sorted unit. Within each session, units with a mean firing rate of at least 0.5 Hz "
  "were retained (1,305 to 2,395 units per session). No units were excluded on the basis of "
  "expected signal, and no area or cell-type pre-selection was applied; area and cell type "
  "are treated as potential findings, not inputs.")
H("Correlation architecture and its comparison", 12)
P("Spike trains were binned at 100 ms. Within each visual stimulus condition (e.g., drifting "
  "gratings, static gratings, natural scenes, gabor patches, flashes, spontaneous activity), "
  "we computed the full unit-by-unit Spearman correlation matrix. Architectures were compared "
  "across all condition pairs within a session (2,052 condition pairs across 30 sessions). "
  "Overall reorganization between two conditions was quantified with the unsigned weighted "
  "Jaccard index (a weighted generalization of the Jaccard 1912 set-overlap coefficient) on the "
  "absolute correlation entries; reorganization is reported as 1 minus this index. Random seed 42 was used throughout; all condition correlation matrices were "
  "saved to disk for reproducibility.")
H("Direct, magnitude-stratified sign-flip rate", 12)
P("To measure correlation sign reversal directly, for every neuron pair we tested whether the "
  "sign of its correlation differed between conditions (sign(r_A) != sign(r_B)). Because a sign "
  "reversal is only architecturally meaningful when both correlations have real magnitude, the "
  "rate was stratified by min(|r_A|, |r_B|) at thresholds 0.0, 0.05, 0.10, 0.15, 0.20 and 0.30, "
  "and each stratum was compared against a within-condition split-half null: each condition's "
  "samples were split into two independent halves, a correlation matrix was recomputed on each "
  "half, and the sign-flip rate was measured between two samples of the same condition (the "
  "reversal rate expected from sampling alone), with sessions as the replication unit (Wilcoxon "
  "across sessions per stratum). The 50% chance rate is not an adequate null, because strong "
  "correlations reflecting shared input are not expected to reverse at coin-flip rates between any "
  "two conditions; the within-condition null is the appropriate comparison. We do not use "
  "any gap between signed-handling and unsigned similarity as a sign-inversion measure, because "
  "such gaps conflate magnitude and sign change; the direct, stratified rate is reported instead.")
H("Per-unit decomposition and waveform association", 12)
P("For each neuron we computed, between each condition pair, its row reorganization (1 minus the "
  "weighted Jaccard of its coupling profile, self excluded) and its sign coherence (the fraction "
  "of its partners preserving correlation sign), the latter across the same swept magnitude "
  "floors. We summarized neurons by the joint of these two continuous quantities and, descriptively, "
  "by four regions of that plane (stable-hub, coherent-magnitude, split-learner, incoherent). To "
  "ask whether spike type relates to reorganization without imposing a cell-type label, we "
  "correlated each neuron's raw extracellular waveform features (duration, halfwidth, spread, "
  "peak-trough ratio, repolarization and recovery slopes) with its reorganization and coherence "
  "using Spearman correlation computed within each session and aggregated across the 30 sessions "
  "(Wilcoxon signed-rank on the per-session correlations). No clustering or trained classifier was "
  "used; categories were not imposed.")
H("Scalar baseline and coupling regimes", 12)
P("As a baseline we computed the field-standard scalar (mean absolute correlation) per condition "
  "and compared its change between conditions to the architecture reorganization. To estimate how "
  "many distinct coupling patterns a neuron expresses across conditions, we computed the "
  "participation ratio of its coupling profile across the condition set (a threshold-free effective "
  "dimensionality; ~1 indicates a single stable pattern). Observed participation ratios were "
  "compared to a condition-shuffled null. Inference throughout treats sessions, not individual "
  "non-independent pairs, as the replication unit where appropriate.")

doc.add_page_break()
H("RESULTS")
H("Reorganization is predominantly in magnitude, with a small but genuine sign component", 12)
P("Across 2,052 condition pairs, correlation architecture reorganized substantially between "
  "visual conditions (median unsigned reorganization 0.58). Sign reversal was rare and fell "
  "monotonically with correlation magnitude: the direct, magnitude-stratified sign-flip rate was "
  "38.8% across all pairs and declined to 4.3% at min(|r|) >= 0.10, 1.2% at >= 0.20, and 0.5% at "
  ">= 0.30, so at meaningful correlation magnitude the reorganization is predominantly a "
  "rearrangement of correlation strength. These rates are not, however, sampling artifacts. "
  "Tested against a within-condition split-half null (the sign-flip rate expected from splitting a "
  "single condition into two independent samples), with sessions as the replication unit, the "
  "observed between-condition rate significantly exceeded the sampling floor at every stratum "
  "(all 30 sessions, Wilcoxon p < 0.001): 38.8% observed vs 23.4% null at all pairs, 4.3% vs 0.9% "
  "at |r| >= 0.10, and 1.2% vs 0.01% at |r| >= 0.20. The excess over the sampling floor therefore "
  "grows with correlation magnitude (from ~1.7-fold to ~100-fold), meaning that the sign reversals "
  "occurring among strongly correlated pairs are almost entirely genuine reorganization rather than "
  "noise. Sign reversal is thus a small but real and magnitude-scaling component of the "
  "reorganization. This comparison is conservative by construction: the within-condition null uses "
  "half the samples per condition, so its correlation matrices are noisier and it overestimates the "
  "true sampling floor; the observed rate therefore clears an inflated bar, and the genuine sign "
  "component is if anything underestimated. The earlier 50% chance rate is not an adequate null and "
  "is not used.")
H("Most neurons restructure magnitude while preserving sign", 12)
P("The per-unit decomposition, binning the continuous reorganization-coherence plane into four "
  "descriptive regions (not discrete classes), showed that 56% of neurons restructured their coupling "
  "profile in strength while preserving sign (coherent-magnitude), 28% changed little (stable-hub), "
  "about 4.2% split their partners by sign (split-learner), and essentially none were unstructured "
  "(~0% incoherent). Row reorganization and sign coherence traded off continuously and consistently: "
  "neurons that restructured more preserved sign less (Spearman rho = -0.52, the same sign in all 30 "
  "of 30 sessions). This trade-off was not solely an artifact of correlations passing through zero: "
  "restricting sign coherence to partners with correlation magnitude at least 0.2 in both conditions "
  "weakened but did not abolish it (rho = -0.21, same sign in 29 of 30 sessions). The split-learner "
  "neurons are therefore the tail of a continuous trade-off, not a discrete category.")
H("Reorganization is only weakly related to spike waveform", 12)
P("Computed within each session and aggregated across sessions, a neuron's reorganization and sign "
  "coherence were related to its extracellular waveform shape consistently but weakly: the largest "
  "association was waveform spread (Spearman rho approximately 0.24 with sign coherence, same sign in "
  "29 of 30 sessions), and the cell-type-relevant waveform halfwidth reached rho approximately 0.16. "
  "Broader, slower waveforms preserved sign somewhat more; narrower, faster waveforms reorganized "
  "somewhat more. The effects are small, indicating that reorganization is a graded property of the "
  "whole population rather than the signature of a distinguishable cell class.")
H("A scalar summary discards most of the reorganization", 12)
P("Between the same condition pairs, the field-standard scalar (mean absolute correlation) changed by "
  "a median of 0.009 while the architecture reorganized by a median of 0.58, roughly a 60-fold "
  "difference; the scalar tracked the direction of reorganization (rho = 0.58) but compressed its "
  "magnitude almost entirely. The architecture-level view therefore resolves a large reorganization "
  "that a scalar summary flattens.")
H("Neurons express a small, structured set of coupling regimes", 12)
P("Across conditions, each neuron's coupling reconfiguration was structured rather than arbitrary. The "
  "participation ratio of a neuron's coupling profile across conditions (an effective count of distinct "
  "coupling patterns) had a median of approximately 4 of a possible ~12, far below a condition-shuffled "
  "null (median approximately 11); every neuron fell below the null median. We emphasize that this "
  "participation ratio is computed over a single neuron's coupling profile across conditions, not over "
  "the eigenspectrum of population activity within a condition; it is therefore a statement about how "
  "structured each neuron's cross-condition reconfiguration is, and not a claim about the dimensionality "
  "of the population code, which is high and may scale with neuron number (Stringer et al. 2019; Manley "
  "et al. 2024). Neurons reconfigure their coupling along a constrained, non-random set of patterns "
  "across visual conditions.")

doc.add_page_break()
H("DISCUSSION")
for t in [
 "Across 30 Neuropixels recordings, the pairwise correlation architecture of mouse visual cortex "
 "reorganizes substantially between visual conditions. That reorganization is predominantly a change in "
 "correlation strength: sign reversal is rare and falls to about 1% of pairs once a modest correlation "
 "magnitude is required. It is not, however, merely sampling noise. When the observed between-condition "
 "sign-flip rate is compared to a within-condition split-half null, it significantly exceeds the "
 "sampling floor at every magnitude stratum (all 30 sessions), and the excess grows with magnitude, so "
 "the sign reversals that do occur among strongly correlated pairs are genuine reorganization. The "
 "honest summary is therefore magnitude-dominated reorganization with a small but real, magnitude-"
 "scaling sign component, not sign preservation. The population and per-unit views agree: most neurons "
 "restructure coupling magnitude while largely preserving sign, reorganization and sign coherence trade "
 "off continuously, and a small minority split their partners by sign.",
 "This characterization depends on two methodological choices: measuring sign reversal directly and "
 "stratifying by magnitude, and testing it against a structure-preserving (within-condition) null "
 "rather than a chance baseline. Unsigned similarity discards sign entirely, and inferring a sign-"
 "inversion fraction from the gap between a sign-handling and an unsigned similarity overstates sign "
 "change because such gaps are dominated by magnitude changes among weak correlations. Equally, a 50% "
 "chance baseline is the wrong null: strong correlations reflect shared input and do not flip at coin-"
 "flip rates, so comparing to 50% makes a trivial observation look like a finding. Only the direct "
 "rate against the within-condition null separates the small genuine sign component from the sampling "
 "floor.",
 "The reorganization is only weakly related to extracellular waveform shape, so it is a graded property "
 "of the population rather than a property carried by a distinct cell class. We deliberately avoided "
 "imposing an excitatory/inhibitory or fast-spiking/regular-spiking label, both because the extracellular "
 "waveform cannot cleanly separate the relevant interneuron classes and because the data show a continuous "
 "gradient rather than discrete groups. The structured per-neuron coupling reconfiguration (participation "
 "ratio far below the condition-shuffled null) indicates that each neuron reconfigures along a "
 "constrained, non-random set of coupling patterns across conditions. This is a statement about the "
 "constraint on a single neuron's cross-condition coupling, computed over its coupling profile, and is "
 "distinct from the dimensionality of the population code measured by the eigenspectrum of population "
 "activity within a condition, which is high-dimensional and may scale without bound (Stringer et al. "
 "2019; Cunningham and Yu 2014; Manley et al. 2024). We interpret the constrained reconfiguration as a "
 "relational description of a neuron's coupling structure, not as a claim about its anatomical inputs.",
 "A magnitude-dominated reorganization signature has also been observed in other, non-neural "
 "correlation systems. We note the convergence as context but do not test it here; the present claims "
 "are specific to visual cortical population recordings.",
]: P(t)
H("Limitations", 12)
P("The analysis is correlational and does not address mechanism or causation. Spike-sorting attribution "
  "and the firing-rate criterion are the only analyst-imposed steps and were held fixed; prior work on "
  "this dataset reports that the population correlation results are robust to spike-sorting quality, "
  "running speed, and bin size. The fundamental unit is the somatic spike train, which integrates a "
  "neuron's dendritic inputs; sign reversals that cancel before the soma are therefore not observable, so "
  "the reported sign-flip rate is a conservative lower bound on sub-somatic reorganization. The per-neuron "
  "coupling-regime count is bounded by the number of conditions and is a relational, not anatomical, "
  "measure. The waveform associations, while consistent across sessions, are small and should not be over-"
  "interpreted as cell-type effects.")
H("Conclusions", 12)
P("Visual cortical population correlation architecture is reorganized between stimulus conditions "
  "predominantly through changes in correlation magnitude, with a small but genuine component of sign "
  "reversal that exceeds the within-condition sampling floor at every magnitude and grows in relative "
  "size with correlation strength. The reorganization is a graded property of the population that a "
  "scalar summary does not resolve, and its sign component is visible only to a direct, magnitude-"
  "stratified measurement tested against a structure-preserving null.")

doc.add_page_break()
H("DISCLOSURES")
P("The author declares no competing financial interests. The author is the founder of Inner Architecture LLC.")
H("AUTHOR CONTRIBUTIONS")
P("D.H.H. conceived and designed the analysis, performed the analysis, and wrote the manuscript.")
H("DECLARATION OF GENERATIVE AI USE")
P("Claude (Anthropic) was used as a programming assistant for pipeline development, code review, and "
  "manuscript formatting. All analytical decisions, methodology, interpretation, and conclusions are the "
  "author's. The tool was not used to generate scientific text, interpret results, or formulate hypotheses.")
H("DATA AND CODE AVAILABILITY")
P("Data are from the Allen Brain Observatory Visual Coding Neuropixels dataset (DANDI 000021). All analysis "
  "code is openly available at https://github.com/nwharbert8-ui/neuropixels-correlation-architecture-wj.")

doc.add_page_break()
H("REFERENCES")
for r in [
 "Cohen MR, Kohn A. Measuring and interpreting neuronal correlations. Nat Neurosci 14: 811-819, 2011. "
 "https://doi.org/10.1038/nn.2842",
 "Cunningham JP, Yu BM. Dimensionality reduction for large-scale neural recordings. Nat Neurosci 17: "
 "1500-1509, 2014. https://doi.org/10.1038/nn.3776",
 "Harbert DH. Sigma-1 and sigma-2 receptors exhibit divergent genome-wide co-expression architectures in "
 "human brain despite shared subcellular localization. Front Pharmacol 17: 1830847, 2026. "
 "https://doi.org/10.3389/fphar.2026.1830847",
 "Jaccard P. The distribution of the flora in the alpine zone. New Phytol 11: 37-50, 1912. "
 "https://doi.org/10.1111/j.1469-8137.1912.tb05611.x",
 "Manley J, Lu S, Barber K, et al. Simultaneous, cortex-wide dynamics of up to 1 million neurons reveal "
 "unbounded scaling of dimensionality with neuron number. Neuron 112: 1694-1709, 2024. "
 "https://doi.org/10.1016/j.neuron.2024.02.011",
 "Niell CM, Stryker MP. Highly selective receptive fields in mouse visual cortex. J Neurosci 28: 7520-7536, "
 "2008. https://doi.org/10.1523/JNEUROSCI.0623-08.2008",
 "Siegle JH, Jia X, Durand S, et al. Survey of spiking in the mouse visual system reveals functional "
 "hierarchy. Nature 592: 86-92, 2021. https://doi.org/10.1038/s41586-020-03171-x",
 "Stringer C, Pachitariu M, Steinmetz N, Carandini M, Harris KD. High-dimensional geometry of population "
 "responses in visual cortex. Nature 571: 361-365, 2019. https://doi.org/10.1038/s41586-019-1346-5",
]:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 2.0; p.paragraph_format.space_after = Pt(2)
    p.add_run(r).font.size = Pt(12)

doc.save(os.path.join(OUT, "Manuscript.docx"))
import docx as _d
words = sum(len(p.text.split()) for p in Document(os.path.join(OUT, "Manuscript.docx")).paragraphs)
print("WROTE", os.path.join(OUT, "Manuscript.docx"), "| approx words:", words)
